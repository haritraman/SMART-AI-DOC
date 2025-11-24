from docx import Document

def build_docx(project, sections, filepath: str):
    doc = Document()
    doc.add_heading(project.title, level=1)

    for sec in sections:
        doc.add_heading(sec.title, level=2)
        content = sec.current_content or ""
        paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
        if not paragraphs:
            doc.add_paragraph("")
        else:
            for p in paragraphs:
                doc.add_paragraph(p)

    doc.save(filepath)
    return filepath
